from docx.shared import RGBColor

from src.utils import limpiar_valor

NO_RESPONDIDO = "(No respondido en el formulario)"


def agregar_campo(tabla, nombre, valor):

    fila = tabla.add_row().cells

    fila[0].text = str(nombre)

    valor = limpiar_valor(valor)

    p = fila[1].paragraphs[0]

    if valor == NO_RESPONDIDO:

        run = p.add_run(valor)

        run.bold = True

        run.font.color.rgb = RGBColor(
            255,
            102,
            0
        )

    else:

        p.add_run(str(valor))


def generar_red(doc, row):

    nombre_registro = limpiar_valor(
        row["Nombre Asesoría"]
    )

    doc.add_heading(
        f"Registro Asociado A: {nombre_registro}",
        level=1
    )

    info = doc.add_table(
        rows=1,
        cols=2
    )

    info.style = "Table Grid"

    info.rows[0].cells[0].text = "Campo"
    info.rows[0].cells[1].text = "Valor"

    agregar_campo(
        info,
        "Correo",
        row.get("Correo electrónico", "")
    )

    agregar_campo(
        info,
        "Nombre",
        row.get("Nombre", "")
    )

    agregar_campo(
        info,
        "Supervisor",
        row.get("Supervisor", "")
    )

    agregar_campo(
        info,
        "Tipo Asesoría",
        row.get("Tipo Asesoría", "")
    )

    doc.add_heading(
        "Brechas Territoriales Lectura",
        level=2
    )

    doc.add_paragraph(
        limpiar_valor(
            row.get(
                "Indique las brechas o nudos críticos comunes identificados en lectura (a nivel de red)",
                ""
            )
        )
    )

    doc.add_heading(
        "Brechas Territoriales Matemática",
        level=2
    )

    doc.add_paragraph(
        limpiar_valor(
            row.get(
                "Indique las brechas o nudos críticos comunes identificados en matemática (a nivel de red)",
                ""
            )
        )
    )

    doc.add_heading(
        "Hallazgos DIA De La Red",
        level=2
    )

    doc.add_paragraph(
        limpiar_valor(
            row.get(
                "Principales hallazgos del análisis DIA de la Red (Socioemocional / Académico): ¿Qué patrones comunes se repiten en el territorio?",
                ""
            )
        )
    )

    doc.add_heading(
        "Fortalezas Transferibles",
        level=2
    )

    doc.add_paragraph(
        limpiar_valor(
            row.get(
                "Fortalezas transferibles de la red (prácticas exitosas, recursos o liderazgos de ciertos establecimientos que pueden modelar o compartirse con el resto de la red)",
                ""
            )
        )
    )

    doc.add_heading(
        "Oportunidades De Mejora Colectiva",
        level=2
    )

    doc.add_paragraph(
        limpiar_valor(
            row.get(
                "Oportunidades de mejora colectiva (nudos críticos comunes, debilidades instaladas en el territorio o necesidades de capacitación transversal)",
                ""
            )
        )
    )

    doc.add_heading(
        "Trabajo Colaborativo De La Red",
        level=2
    )

    tabla_red = doc.add_table(
        rows=1,
        cols=2
    )

    tabla_red.style = "Table Grid"

    tabla_red.rows[0].cells[0].text = "Elemento"
    tabla_red.rows[0].cells[1].text = "Valor"

    agregar_campo(
        tabla_red,
        "Lenguaje",
        row.get(
            "Existen acciones articuladas o transversales en los PME de los establecimientos para abordar la asignatura de lenguaje",
            ""
        )
    )

    agregar_campo(
        tabla_red,
        "Acuerdo Lenguaje",
        row.get(
            "Indique una breve descripción del acuerdo de red",
            ""
        )
    )

    agregar_campo(
        tabla_red,
        "Matemática",
        row.get(
            "¿Existen acciones articuladas o transversales en los PME de los establecimientos para abordar la asignatura de matemática?",
            ""
        )
    )

    agregar_campo(
        tabla_red,
        "Acuerdo Matemática",
        row.get(
            "Indique una breve descripción del acuerdo de red2",
            ""
        )
    )

    agregar_campo(
        tabla_red,
        "Ajuste Territorial",
        row.get(
            "Verificación de ajuste territorial: Las metas de los PME ¿apuntan colectivamente a resolver las brechas identificadas en la red?",
            ""
        )
    )

    agregar_campo(
        tabla_red,
        "Sugerencias PME",
        row.get(
            "Sugerencias de ajuste a los PME para potenciar el trabajo colaborativo de la Red",
            ""
        )
    )

    agregar_campo(
        tabla_red,
        "Plan de Trabajo",
        row.get(
            "¿Se establece un plan de trabajo o calendario de nodos para el semestre?",
            ""
        )
    )

    agregar_campo(
        tabla_red,
        "Compromiso Próxima Sesión",
        row.get(
            "Compromiso clave de la red para la próxima sesión",
            ""
        )
    )

    doc.add_page_break()