from docx.shared import RGBColor

from src.utils import limpiar_valor

NO_RESPONDIDO = "(No respondido en el formulario)"


def agregar_campo(tabla, nombre, valor):

    fila = tabla.add_row().cells

    fila[0].text = str(nombre)

    valor = limpiar_valor(valor)

    p = fila[1].paragraphs[0]

    if valor == NO_RESPONDIDO:

        run = p.add_run(valor)

        run.bold = True

        run.font.color.rgb = RGBColor(
            255,
            102,
            0
        )

    else:

        p.add_run(str(valor))


def generar_directa(doc, row):

    nombre_registro = limpiar_valor(
        row["Nombre Asesoría"]
    )

    doc.add_heading(
        f"Registro Asociado A: {nombre_registro}",
        level=1
    )

    info = doc.add_table(
        rows=1,
        cols=2
    )

    info.style = "Table Grid"

    info.rows[0].cells[0].text = "Campo"
    info.rows[0].cells[1].text = "Valor"

    agregar_campo(
        info,
        "Correo",
        row.get("Correo electrónico", "")
    )

    agregar_campo(
        info,
        "Nombre",
        row.get("Nombre", "")
    )

    agregar_campo(
        info,
        "Supervisor",
        row.get("Supervisor", "")
    )

    agregar_campo(
        info,
        "Director",
        row.get("Nombre Director", "")
    )

    agregar_campo(
        info,
        "Tipo Asesoría",
        row.get("Tipo Asesoría", "")
    )

    doc.add_heading(
        "Brechas Críticas Lectura",
        level=2
    )

    doc.add_paragraph(
        limpiar_valor(
            row.get(
                "Indique las brechas críticas identificadas en lectura",
                ""
            )
        )
    )

    doc.add_heading(
        "Brechas Críticas Matemática",
        level=2
    )

    doc.add_paragraph(
        limpiar_valor(
            row.get(
                "Indique las brechas críticas identificadas en matemática",
                ""
            )
        )
    )

    doc.add_heading(
        "Hallazgos DIA",
        level=2
    )

    doc.add_paragraph(
        limpiar_valor(
            row.get(
                "Principales hallazgos del análisis DIA (Socioemocional / Académico)",
                ""
            )
        )
    )

    doc.add_heading(
        "Fortalezas",
        level=2
    )

    doc.add_paragraph(
        limpiar_valor(
            row.get(
                "Fortalezas del Establecimiento (prácticas exitosas, recursos o liderazgos consolidados)",
                ""
            )
        )
    )

    doc.add_heading(
        "Oportunidades De Mejora",
        level=2
    )

    doc.add_paragraph(
        limpiar_valor(
            row.get(
                "Oportunidades de mejora del Establecimiento (nudos críticos o debilidades a subsanar)",
                ""
            )
        )
    )

    doc.add_heading(
        "PME",
        level=2
    )

    pme = doc.add_table(
        rows=1,
        cols=2
    )

    pme.style = "Table Grid"

    pme.rows[0].cells[0].text = "Elemento"
    pme.rows[0].cells[1].text = "Valor"

    agregar_campo(
        pme,
        "Acciones Lenguaje",
        row.get(
            "El PME vigente contempla acciones específicas para el abordaje de la asignatura de lenguaje",
            ""
        )
    )

    agregar_campo(
        pme,
        "Descripción Lenguaje",
        row.get(
            "Indique una breve descripción de la acción",
            ""
        )
    )

    agregar_campo(
        pme,
        "Acciones Matemática",
        row.get(
            "El PME vigente contempla acciones específicas para el abordaje de la asignatura de matemática",
            ""
        )
    )

    agregar_campo(
        pme,
        "Descripción Matemática",
        row.get(
            "Indique una breve descripción de la acción2",
            ""
        )
    )

    agregar_campo(
        pme,
        "Observaciones PME",
        row.get(
            "Observaciones / Sugerencias de ajuste al PME",
            ""
        )
    )

    doc.add_page_break()