from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from src.utils import limpiar_valor
from src.plantillas.directa import generar_directa
from src.plantillas.red import generar_red


def configurar_estilos(doc):

    for nombre in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:

        estilo = doc.styles[nombre]

        estilo.font.name = "Aptos"


def agregar_encabezado(doc):

    try:

        section = doc.sections[0]

        header = section.header

        p = header.paragraphs[0]

        run = p.add_run()

        run.add_picture(
            "assets/ApoyoalaMejora.png",
            width=Inches(1.3)
        )

    except:
        pass


def generar_word(df, salida):

    doc = Document()

    configurar_estilos(doc)

    agregar_encabezado(doc)

    region = limpiar_valor(
        df.iloc[0]["Indique su región"]
    )

    deprov = limpiar_valor(
        df.iloc[0]["Deprov"]
    )

    modalidad = limpiar_valor(
        df.iloc[0]["Tipo Asesoría"]
    )

    asesor = limpiar_valor(
        df.iloc[0]["Nombre"]
    )

    titulo = doc.add_paragraph()

    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = titulo.add_run(
        "Informe Individual Etapa De Preparación De La Asesoría"
    )

    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(18)

    run.font.color.rgb = RGBColor(
        0,
        70,
        140
    )

    doc.add_paragraph()

    tabla_portada = doc.add_table(
        rows=4,
        cols=2
    )

    tabla_portada.style = "Table Grid"

    tabla_portada.cell(0, 0).text = "Región"
    tabla_portada.cell(0, 1).text = region

    tabla_portada.cell(1, 0).text = "DEPROV"
    tabla_portada.cell(1, 1).text = deprov

    tabla_portada.cell(2, 0).text = "Modalidad"
    tabla_portada.cell(2, 1).text = modalidad

    tabla_portada.cell(3, 0).text = "Asesor"
    tabla_portada.cell(3, 1).text = asesor

    doc.add_page_break()

    for _, row in df.iterrows():

        tipo = limpiar_valor(
            row.get("Tipo Asesoría", "")
        )

        if tipo == "Directa EE":

            generar_directa(
                doc,
                row
            )

        elif tipo == "Red EE":

            generar_red(
                doc,
                row
            )

        else:

            generar_directa(
                doc,
                row
            )

    doc.save(salida)